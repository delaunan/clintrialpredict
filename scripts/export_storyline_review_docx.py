#!/usr/bin/env python
"""Convert the storyline Markdown review pack into a formatted Word document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


DEFAULT_MARKDOWN = Path("reports/narrative_evals/final_validation_storyline_review_pack.md")
DEFAULT_OUT = Path("reports/narrative_evals/final_validation_storyline_review_pack.docx")

LABEL_LINES = {
    "Reviewer notes:",
    "Consistency note:",
    "Completion Outlook narrative:",
    "Design Confidence narrative:",
    "Design Confidence subcategory ratings and rationales:",
    "Questions:",
    "Findings:",
    "Changed fields:",
}


def _clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("\\|", "|")
    return text.strip()


def _set_default_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in (
        ("Heading 1", 15),
        ("Heading 2", 12),
        ("Heading 3", 10.5),
    ):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(9.5)


def _set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def _add_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(_clean_inline(text))
    run.bold = True


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index].cells
        for col_index, value in enumerate(row):
            cell = cells[col_index]
            cell.text = _clean_inline(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(7.5 if len(row) > 4 else 8.5)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph()


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1

    rows: list[list[str]] = []
    for offset, line in enumerate(table_lines):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if offset == 1 and all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def _add_paragraph(document: Document, text: str) -> None:
    cleaned = _clean_inline(text)
    if not cleaned:
        return
    document.add_paragraph(cleaned)


def convert_markdown_to_docx(markdown_path: Path, out_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    _set_default_styles(document)
    _set_margins(document)

    index = 0
    first_trial_heading = True
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            document.add_heading(_clean_inline(line[2:]), level=0)
            index += 1
            continue

        if line.startswith("## "):
            if first_trial_heading:
                first_trial_heading = False
            else:
                document.add_page_break()
            document.add_heading(_clean_inline(line[3:]), level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(_clean_inline(line[4:]), level=2)
            index += 1
            continue

        if line.startswith("|"):
            rows, index = _parse_table(lines, index)
            _add_table(document, rows)
            continue

        if line in LABEL_LINES:
            _add_label(document, line)
            index += 1
            continue

        if line.startswith("- "):
            document.add_paragraph(_clean_inline(line[2:]), style="List Bullet")
            index += 1
            continue

        if line.startswith(">"):
            paragraph = document.add_paragraph(_clean_inline(line.lstrip("> ")))
            paragraph.style = "Intense Quote"
            index += 1
            continue

        # Merge adjacent plain lines into a single paragraph until the next structural marker.
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith("- ")
                or next_line.startswith(">")
                or next_line in LABEL_LINES
            ):
                break
            paragraph_lines.append(next_line)
            index += 1
        _add_paragraph(document, " ".join(paragraph_lines))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MARKDOWN)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    convert_markdown_to_docx(args.markdown, args.out)
    print(f"Wrote {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
